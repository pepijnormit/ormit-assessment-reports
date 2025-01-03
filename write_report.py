import os
import sys
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import ast
import json
import re

print('YOU NEED TO COMMENT THIS OUT BEFORE EXE')
script_directory = os.path.dirname(os.path.abspath(__file__))
os.chdir(script_directory)  

def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except AttributeError:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

### Initial clean:
def clean(text):
    if not text or isinstance(text, list):  # If empty or already a list
        return text
    elif isinstance(text, str):  # If string
        # Remove any square bracket items and return their contents (cleaning some classic GPT style aspects)
        cleaned_text = re.sub(r'\【.*?\】', '', text).strip()
        cleaned_text = cleaned_text.replace("```python", "")
        cleaned_text = cleaned_text.replace("```", "")
        cleaned_text = cleaned_text.replace("**", "")
        cleaned_text = strip_extra_quotes(cleaned_text)
        return cleaned_text  # Return cleaned string immediately if no lists to extract

def strip_extra_quotes(input_string):
    # Check if the string starts and ends with double quotes
    if input_string.startswith('"') and input_string.endswith('"'):
        # Remove the extra quotes
        return input_string[1:-1]
    return input_string

#######################  Actual Work  ##########################
def clean_up(loc_dic):
    with open(loc_dic, 'r') as json_file:
        loaded_data = json.load(json_file)
    results_list = {}
    for key in loaded_data.keys():
        results_list[key] = clean(loaded_data[key]) #Extract list from string
    return results_list

### Put it in Word:
#Icons
template = resource_path('resources/template.docx')
image_path_improv = resource_path("resources/improvement.png")
image_path_average = resource_path("resources/average.png")
image_path_strong = resource_path("resources/strong.png")

def update_document(output_dic, name, assessor, gender):
    doc = Document(template)
    ### CONTENT
    #Personal details
    add_content_detailstable(doc, [name, "", "MCP", "", ""])
    replace_and_format_header_text(doc, name)
    replace_placeholder_in_docx(doc, '***', name.split()[0], font_name='Montserrat Light') 
    
    #Add assessor
    replace_placeholder_in_docx(doc, 'ASSESSOR', assessor, font_name='Montserrat Light') 
        
    #First impression
    if "prompt2_firstimpr" in output_dic and output_dic['prompt2_firstimpr'] != "":
        firstimpr_pietless = replacePiet(output_dic['prompt2_firstimpr'], name, gender)
        add_content_below_heading(doc, "First impression", firstimpr_pietless, "First impression")
    
    #Personality
    if "prompt3_personality" in output_dic and output_dic['prompt3_personality'] != "":
        personality_pietless = replacePiet(output_dic['prompt3_personality'], name, gender)
        add_content_below_heading(doc, "Personality", personality_pietless, "Personality")
   
    #Cognitive Capacitiy Test results
    if "prompt4_cogcap" in output_dic and output_dic['prompt4_cogcap'] != "":
        add_content_cogcaptable(doc, output_dic['prompt4_cogcap'])
   
    #Add language levels
    if "prompt5_language" in output_dic and output_dic['prompt5_language'] != "":
        language_skills(doc, output_dic['prompt5_language'])
    
    #Add conclusion columns
    if "prompt6a_conqual" in output_dic and output_dic['prompt6a_conqual'] != "":
        conclusion(doc, 0, output_dic['prompt6a_conqual'])
    if "prompt6b_conimprov" in output_dic and output_dic['prompt6b_conimprov'] != "":
        conclusion(doc, 1, output_dic['prompt6b_conimprov'])
    
    # #Profile review
    if 'prompt7_qualscore' in output_dic and output_dic['prompt7_qualscore'] != '' and 'prompt7_improvscore' in output_dic and output_dic['prompt7_improvscore'] != '':
        add_icons2(doc, output_dic['prompt7_qualscore'], output_dic['prompt7_improvscore'])

    # Save the document
    current_time = datetime.now()
    formatted_time = current_time.strftime("%m%d%H%M")  # Format as MMDDHHMinMin
    updated_doc_path = f"Assessment Report - {name} - {formatted_time}.docx"
    doc.save(updated_doc_path)
    os.startfile(updated_doc_path)
    return updated_doc_path

def replacePiet(text, name, gender):
    text = text.replace("Piet", name.split()[0])
    text = re.sub(r'\bthe trainee\b', name.split()[0], text, flags=re.IGNORECASE) #Replace "The trainee"
    
    if gender == 'M':
        replacements = {
            "She": "He",
            "she": "he",
            "Her": "Him",
            "her": "him",
            "Hers": "His",
            "hers": "his",
            "Herself": "Himself",
            "herself": "himself",
        }
        for female, male in replacements.items():
            text = re.sub(r'\b' + re.escape(female) + r'\b', male, text)

    elif gender == 'F':
        replacements = {
            "He": "She",
            "he": "she",
            "Him": "Her",
            "him": "her",
            "His": "Her",
            "his": "her",
            "Himself": "Herself",
            "himself": "herself",
        }
        for male, female in replacements.items():
            text = re.sub(r'\b' + re.escape(male) + r'\b', female, text)

    elif gender == 'X':
        replacements = {
            "He": "They",
            "he": "they",
            "Him": "Them",
            "him": "them",
            "His": "Their",
            "his": "their",
            "Himself": "Themselves",
            "himself": "themselves",
            "She": "They",
            "she": "they",
            "Her": "Their",
            "her": "their",
            "Herself": "Themself",
            "herself": "themself",
        } 
        for prev, nonbi in replacements.items():
            text = re.sub(r'\b' + re.escape(prev) + r'\b', nonbi, text)
    
    return text
    
def restructure_date(date_str):
    # Replace / with - to handle both delimiters
    date_str = date_str.replace('/', '-')
    
    try:
        # Try parsing the date in DD-MM-YYYY format
        datetime.strptime(date_str, '%d-%m-%Y')
        return date_str  # Return as is if it's already in correct format
    except ValueError:
        try:
            # Parse the input date string in YYYY-MM-DD format
            date_obj = datetime.strptime(date_str, '%Y-%m-%d')
            # Restructure and return the date in DD-MM-YYYY format
            return date_obj.strftime('%d-%m-%Y')
        except ValueError:
            return ''  # Return empty string if parsing fails

def set_font_properties(cell):
    # Set the font properties directly by modifying the cell's XML
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Montserrat Light'  # Set the font name
            run.font.size = Pt(11)  # Set font size if needed

            # Create and set font properties for Montserrat Light
            r = run._element
            rPr = r.rPr
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                r.append(rPr)

            rFonts = OxmlElement('w:rFonts')
            rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii', 'Montserrat Light')
            rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi', 'Montserrat Light')
            rPr.append(rFonts)

def set_font_properties2(para):
    # Get the full text of the paragraph
    full_text = para.text
    # Clear the existing text runs
    para.clear()  
    
    # Split the text into lines based on new lines
    lines = full_text.splitlines()
    
    for line in lines:
        # Split the line into words
        words = line.split()
        # If there are words in the line
        if words:
            # Create runs for all words except the last one
            for word in words[:-1]:
                run = para.add_run(word + ' ')  # Add space after each word
                run.font.name = 'Montserrat Light'
                run.font.size = Pt(10)
                run.bold = False  # Ensure non-bold for non-last words
                r = run._element
                rPr = r.rPr
                
                if rPr is None:
                    rPr = OxmlElement('w:rPr')
                    r.append(rPr)

                rFonts = OxmlElement('w:rFonts')
                rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii', 'Montserrat Light')
                rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi', 'Montserrat Light')
                rPr.append(rFonts)

            # Add tabs before the last word based on the first word
            if words[0] == 'Dutch':
                para.add_run('\t\t')  # Three tabs for 'Dutch'
            else:
                para.add_run('\t')  # Two tabs for other cases

            last_word = words[-1]
            last_run = para.add_run(last_word)  # Last word without space
            last_run.font.name = 'Montserrat Light'
            last_run.font.size = Pt(10)
            last_run.bold = True  # Bold for last word
            r = last_run._element
            rPr = r.rPr
            
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                r.append(rPr)

            rFonts = OxmlElement('w:rFonts')
            rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii', 'Montserrat Light')
            rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi', 'Montserrat Light')
            rPr.append(rFonts)
        
def add_content_detailstable(doc, personal_details):
    if len(personal_details)==1 and all(isinstance(ele, str) for ele in personal_details): #['allll, infoo, here']
        personal_details = personal_details[0].split(',')
         
    # Access the first table in the document
    table = doc.tables[0]  

    # PAGE 1
    if personal_details != '':
        for row in table.rows:
            # Check the first and second cells in the current row
            if len(row.cells) > 1:  # Ensure there are at least two cells
                first_cell_text = row.cells[0].text.strip()
                second_cell_text = row.cells[1].text.strip()
                            
                # Update for Name candidate
                if first_cell_text == "Name candidate" and second_cell_text == ":":
                    cell = row.cells[2]
                    cell.text = personal_details[0]  # Directly set the text of the cell
                    set_font_properties(cell)
                    # print("Name added successfully.")
    
                # Update for Date of Birth
                if first_cell_text == "Date of birth" and second_cell_text == ":":
                    cell = row.cells[2]
                    cell.text = restructure_date(personal_details[1])
                    set_font_properties(cell)
                    # print("Date of Birth added successfully.")
    
                # Update for Position
                if first_cell_text == "Position" and second_cell_text == ":":
                    cell = row.cells[2]
                    cell.text = personal_details[2]
                    set_font_properties(cell)
                    # print("Position added successfully.")
    
                # Update for Assessment Date
                if first_cell_text == "Assessment date" and second_cell_text == ":":
                    cell = row.cells[2]
                    cell.text = restructure_date(personal_details[3])
                    set_font_properties(cell)
                    # print("Assessment Date added successfully.")
    
                # Update for Pool
                if first_cell_text == "Pool" and second_cell_text == ":":
                    cell = row.cells[2]
                    cell.text = personal_details[4]
                    set_font_properties(cell)
                    # print("Pool added successfully.")
    return

def add_content_below_heading(doc, heading, content, heading_name=None):
    # Split the content into paragraphs
    paragraphs = content.strip().split('\n\n')  # Split by double newline for new paragraphs

    # Find the bold heading and add content below it
    for paragraph in doc.paragraphs:
        if heading in paragraph.text:
            # Check if any run in the paragraph is bold
            for run in paragraph.runs:
                if run.bold:  # Check if the run is bold
                    # Collect the new paragraphs
                    new_paragraphs = []
                    for index, para in enumerate(paragraphs):
                        # Prepend a tab character for all except the first paragraph
                        if index > 0:
                            para = '\t' + para.strip()
                        else:
                            para = para.strip()

                        # Create a new paragraph for each segment of content
                        new_paragraphs.append(para)

                    # Insert all new paragraphs at once, in reverse order
                    for new_para in reversed(new_paragraphs):  # Reverse the order for insertion
                        # Create a new paragraph
                        new_paragraph = doc.add_paragraph(new_para)
                        # Insert the new paragraph at the correct position (after the heading)
                        doc._element.body.insert(doc._element.body.index(paragraph._element) + 1, new_paragraph._element)

                        run = new_paragraph.runs[0]
                        run.font.name = 'Montserrat Light'  # Set the font name
                        run.font.size = Pt(10)  # Set font size

                        # Create and set font properties for Montserrat Light
                        r = run._element
                        rPr = r.rPr
                        if rPr is None:
                            rPr = OxmlElement('w:rPr')
                            r.append(rPr)

                        rFonts = OxmlElement('w:rFonts')
                        rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii', 'Montserrat Light')
                        rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi', 'Montserrat Light')
                        rPr.append(rFonts)

                    return  # Exit after adding the paragraphs
    if heading_name:
        print(f"No bold '{heading_name}' found.")


def add_content_cogcaptable(doc, scores=[]):
    if isinstance(scores, str):
        scores = ast.literal_eval(scores)
    
    if isinstance(scores, list):
        table = doc.tables[1]  # Access the second table
        
        for i in range(6):
            cell = table.rows[1].cells[i + 1]  
            if i == 0:
                paragraph = cell.add_paragraph()
                run = paragraph.add_run(str(scores[i]))  # Convert score to string
                run.bold = True  # Set the run to bold
                run.underline = True  # Set the run to underlined
                run.font.name = 'Montserrat Light'  # Set the font name
                run.font.size = Pt(11)  # Set font size
                paragraph.alignment = 1  # Center alignment
            else:
                cell.text = str(scores[i])  # For other cells, set text normally
                set_font_properties(cell)  # Update font properties for the cell
                for paragraph in cell.paragraphs:
                    paragraph.alignment = 1  # Center alignment

def language_skills(doc, replacements):
    if isinstance(replacements, str) and replacements!="":
        if not re.search(r'["\']', replacements):  # If there are no quotes at all
            # Add quotes around elements that are not already quoted
            replacements = re.sub(r'(\w+)', r'"\1"', replacements)
        replacements = ast.literal_eval(replacements)

    # Define the keywords to identify the section
    section_title = "Language Skills"
    language_names = ["Dutch", "French", "English"]

    # Flag to determine if we are in the relevant section
    in_section = False

    for para in doc.paragraphs:
        if section_title in para.text:
            in_section = True
        elif any(name in para.text for name in language_names):
            if in_section:
                # Iterate over the language names to replace the placeholder
                for name in language_names:
                    if name in para.text and ".." in para.text:
                        # Get the corresponding replacement value
                        replacement_value = replacements.pop(0) if replacements else ""
                        para.text = para.text.replace("..", replacement_value)

                        # Apply font properties directly to the paragraph
                        set_font_properties2(para)  # Ensure the new text is formatted
                        break
        # Stop searching when we exit the section
        elif in_section and para.text.strip() == "":
            break
        
# def add_icons(doc, table_no, res):
#     if isinstance(res, str):
#         res = ast.literal_eval(res)
    
#     if isinstance(res, list):
#         # Access the specified properties table in the document
#         table = doc.tables[table_no]
#         for i in range(4):
#             cell = table.rows[i + 1].cells[0]
            
#             # Clear existing paragraphs in the cell
#             for paragraph in cell.paragraphs:
#                 paragraph.clear()  # Remove any existing text or paragraphs
#                 #Remove whitespace above picture:
#                 p = paragraph._p  # ---this is the paragraph XML element---
#                 p.getparent().remove(p)
#             # Directly add the image without creating a new paragraph
#             if res[i] == 1 or res[i] =='1':
#                 run = cell.add_paragraph().add_run()  # Create a run in a new paragraph
#                 run.add_picture(image_path_improv, width=Inches(.5))
#             elif res[i] == 2 or res[i] =='2':
#                 run = cell.add_paragraph().add_run()  # Create a run in a new paragraph
#                 run.add_picture(image_path_average, width=Inches(.5))
#             elif res[i] == 3 or res[i] =='3':
#                 run = cell.add_paragraph().add_run()  # Create a run in a new paragraph
#                 run.add_picture(image_path_strong, width=Inches(.5))
#     return

def add_icons2(doc, list_plus, list_min, items_per_table=4):
    if isinstance(list_plus, str):
        list_plus = ast.literal_eval(list_plus)
    if isinstance(list_min, str):
        list_min = ast.literal_eval(list_min)
        
    if isinstance(list_plus, list) and isinstance(list_min, list):
        list_res = [i + j for i, j in zip(list_plus, list_min)] #Add per item
        
        for i in range(len(list_res)):  
            table_no = (i // items_per_table) + 4  # Integer division will give you the table number, 4 is the table_no of the first icon table
            row_no = (i % items_per_table)  # This gives row 1-4
            table = doc.tables[table_no]
            
            cell = table.rows[row_no + 1].cells[0]
            # Clear existing paragraphs in the cell
            for paragraph in cell.paragraphs:
                paragraph.clear()  # Remove any existing text or paragraphs
                #Remove whitespace above picture:
                p = paragraph._p  # ---this is the paragraph XML element---
                p.getparent().remove(p)
            # Directly add the image without creating a new paragraph
            if list_res[i] == -1:
                run = cell.add_paragraph().add_run()  # Create a run in a new paragraph
                run.add_picture(image_path_improv, width=Inches(.5))
            elif list_res[i] == 0:
                run = cell.add_paragraph().add_run()  # Create a run in a new paragraph
                run.add_picture(image_path_average, width=Inches(.5))
            elif list_res[i] == 1:
                run = cell.add_paragraph().add_run()  # Create a run in a new paragraph
                run.add_picture(image_path_strong, width=Inches(.5))
    return

def conclusion(doc, column, list_items):  
    if isinstance(list_items, str):
        list_items = ast.literal_eval(list_items)

    # Access the first properties table in the document
    table = doc.tables[2]  
    cell = table.rows[1].cells[column] 
    
    # Clear existing content in the cell
    for paragraph in cell.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)

    # Add bullet points
    for point in list_items:
        paragraph = cell.add_paragraph()          
        text_run = paragraph.add_run(f'\t -{point}')  # Add the text after the bullet
        text_run.font.name = 'Montserrat Light'
        text_run.font.size = Pt(10)

#Last style improvements
def replace_and_format_header_text(doc, new_text):
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            if '***' in paragraph.text:
                # Replace the old text with new text
                paragraph.text = paragraph.text.replace('***', new_text)
                
                # Apply font formatting to each run in the paragraph
                for run in paragraph.runs:
                    run.font.name = 'Montserrat SemiBold'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(*(0xED, 0x6B, 0x55))  # Set the color
                    run.bold = True  # SemiBold approximation
                    run.italic = False  # Ensure the text is not italicized
                    
                    # Ensure compatibility with Montserrat SemiBold through XML handling
                    rFonts = OxmlElement('w:rFonts')
                    rFonts.set(qn('w:ascii'), 'Montserrat SemiBold')
                    rFonts.set(qn('w:hAnsi'), 'Montserrat SemiBold')
                    run._element.rPr.append(rFonts)

def replace_placeholder_in_docx(doc, placeholder, replacement, font_name='Montserrat', font_size=10):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            # Replace the placeholder text
            inline = paragraph.runs
            for i in range(len(inline)):
                if placeholder in inline[i].text:
                    # Replace text and set font properties
                    inline[i].text = inline[i].text.replace(placeholder, replacement)
                    inline[i].font.name = font_name
                    inline[i].font.size = Pt(font_size)
                    
def open_file(file_path):
    # Open the file automatically based on the operating system
    if os.name == 'nt':  # For Windows
        os.startfile(file_path)
    elif os.name == 'posix':  # For macOS and Linux
        os.system(f'open "{file_path}"')